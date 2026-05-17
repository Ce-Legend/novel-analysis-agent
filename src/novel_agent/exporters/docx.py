from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from .report import ReportBlock


def export_docx(report_blocks: list[ReportBlock], output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    _configure_docx_styles(document)

    for block in report_blocks:
        if block.kind == "heading":
            paragraph = document.add_paragraph(style=_heading_style_name(block))
            paragraph.add_run(block.text)
        elif block.kind == "bullet":
            paragraph = _add_bullet_paragraph(document, block)
            if block.style in {"unit_label", "section_item"}:
                _apply_label_runs(paragraph, block.text)
            else:
                paragraph.add_run(block.text)
        else:
            document.add_paragraph(block.text, style="Delivery Body")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _configure_docx_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)

    _ensure_style(styles, "Delivery Title", WD_STYLE_TYPE.PARAGRAPH, base="Title")
    title = styles["Delivery Title"]
    title.font.name = "PingFang SC Semibold"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC Semibold")
    title.font.size = Pt(22)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)

    _ensure_style(styles, "Delivery Heading 1", WD_STYLE_TYPE.PARAGRAPH, base="Heading 1")
    heading1 = styles["Delivery Heading 1"]
    heading1.font.name = "PingFang SC Semibold"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC Semibold")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(8)

    _ensure_style(styles, "Delivery Heading 2", WD_STYLE_TYPE.PARAGRAPH, base="Heading 2")
    heading2 = styles["Delivery Heading 2"]
    heading2.font.name = "PingFang SC Semibold"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC Semibold")
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(6)

    _ensure_style(styles, "Delivery Body", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    body = styles["Delivery Body"]
    body.font.name = "PingFang SC"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    body.font.size = Pt(10.5)
    body.paragraph_format.line_spacing = 1.35
    body.paragraph_format.space_after = Pt(6)

    _ensure_style(styles, "Delivery Bullet", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    bullet = styles["Delivery Bullet"]
    bullet.font.name = "PingFang SC"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Mm(4)
    bullet.paragraph_format.first_line_indent = Mm(-4)
    bullet.paragraph_format.space_after = Pt(5)
    bullet.paragraph_format.line_spacing = 1.3

    _ensure_style(styles, "Delivery Bullet Detail", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    detail = styles["Delivery Bullet Detail"]
    detail.font.name = "PingFang SC"
    detail._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    detail.font.size = Pt(10)
    detail.paragraph_format.left_indent = Mm(10)
    detail.paragraph_format.first_line_indent = Mm(-4)
    detail.paragraph_format.space_after = Pt(4)
    detail.paragraph_format.line_spacing = 1.25


def _ensure_style(styles, name: str, style_type, *, base: str):  # noqa: ANN001
    if name in styles:
        return styles[name]
    style = styles.add_style(name, style_type)
    style.base_style = styles[base]
    return style


def _heading_style_name(block: ReportBlock) -> str:
    if block.level == 1:
        return "Delivery Title"
    if block.level == 2:
        return "Delivery Heading 1"
    return "Delivery Heading 2"


def _add_bullet_paragraph(document: Document, block: ReportBlock):
    style_name = "Delivery Bullet Detail" if block.level >= 2 else "Delivery Bullet"
    paragraph = document.add_paragraph(style=style_name)
    if block.level >= 2:
        paragraph.add_run("• ")
    elif block.style not in {"unit_label", "group_label"}:
        paragraph.add_run("• ")
    if block.style in {"unit_label", "group_label"}:
        _shade_paragraph(paragraph, "F4F0E8")
    return paragraph


def _apply_label_runs(paragraph, text: str) -> None:  # noqa: ANN001
    label, separator, value = text.partition("：")
    if separator:
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        value_run = paragraph.add_run(value)
        value_run.bold = False
    else:
        run = paragraph.add_run(text)
        run.bold = True


def _shade_paragraph(paragraph, fill: str) -> None:  # noqa: ANN001
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
